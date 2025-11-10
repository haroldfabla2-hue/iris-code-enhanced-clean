"""
Herramienta para procesamiento de archivos de documentos.
Soporta PDF, DOCX, TXT y otros formatos comunes.
"""

import os
import mimetypes
import hashlib
import time
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import json

# Imports condicionales para diferentes tipos de archivos
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet
    ENCODING_DETECTION_AVAILABLE = True
except ImportError:
    ENCODING_DETECTION_AVAILABLE = False

from .base_tool import BaseTool, ToolResult

class FileProcessor(BaseTool):
    """Herramienta para procesamiento seguro de archivos de documentos"""
    
    def __init__(self):
        super().__init__(
            name="file_processor",
            description="Procesador de archivos de documentos (PDF, DOCX, TXT) con funcionalidades de seguridad",
            timeout=120
        )
        
        # Tipos MIME permitidos
        self.allowed_mime_types = [
            'text/plain',
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/csv',
            'application/json',
            'text/html',
            'text/markdown'
        ]
        
        # Extensiones permitidas
        self.allowed_extensions = [
            '.txt', '.pdf', '.docx', '.csv', '.json', '.html', '.md', '.log'
        ]
        
        # Tamaño máximo de archivo (10MB)
        self.max_file_size = 10 * 1024 * 1024
        
        # Directorio temporal seguro
        self.temp_dir = "/tmp/file_processor"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def validate_file(self, file_path: str) -> ToolResult:
        """
        Valida que un archivo sea seguro y accesible
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            ToolResult con validación
        """
        self.set_running()
        
        try:
            # Validar ruta
            if not isinstance(file_path, str):
                return self.create_result(
                    success=False,
                    error="Ruta de archivo debe ser string"
                )
            
            # Sanitizar ruta
            sanitized_path = self.sanitize_input(file_path, max_length=500)
            
            # Verificar que el archivo existe
            if not os.path.exists(sanitized_path):
                return self.create_result(
                    success=False,
                    error=f"Archivo no encontrado: {sanitized_path}"
                )
            
            # Verificar que es un archivo (no directorio)
            if not os.path.isfile(sanitized_path):
                return self.create_result(
                    success=False,
                    error=f"La ruta no es un archivo: {sanitized_path}"
                )
            
            # Verificar tamaño
            file_size = os.path.getsize(sanitized_path)
            if file_size > self.max_file_size:
                return self.create_result(
                    success=False,
                    error=f"Archivo demasiado grande: {file_size} bytes (máximo: {self.max_file_size})"
                )
            
            # Verificar tipo MIME
            mime_type, _ = mimetypes.guess_type(sanitized_path)
            if mime_type not in self.allowed_mime_types:
                return self.create_result(
                    success=False,
                    error=f"Tipo MIME no permitido: {mime_type}"
                )
            
            # Verificar extensión
            file_ext = Path(sanitized_path).suffix.lower()
            if file_ext not in self.allowed_extensions:
                return self.create_result(
                    success=False,
                    error=f"Extensión no permitida: {file_ext}"
                )
            
            # Verificar permisos de lectura
            if not os.access(sanitized_path, os.R_OK):
                return self.create_result(
                    success=False,
                    error="Sin permisos de lectura del archivo"
                )
            
            # Calcular hash del archivo
            file_hash = self._calculate_file_hash(sanitized_path)
            
            self.set_completed()
            
            return self.create_result(
                success=True,
                data={
                    'path': sanitized_path,
                    'size': file_size,
                    'mime_type': mime_type,
                    'extension': file_ext,
                    'hash': file_hash,
                    'valid': True
                }
            )
            
        except Exception as e:
            return self.handle_exception(e)
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calcula hash SHA-256 del archivo"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def read_text_file(self, file_path: str, encoding: Optional[str] = None) -> ToolResult:
        """
        Lee un archivo de texto de forma segura
        
        Args:
            file_path: Ruta al archivo
            encoding: Codificación específica (opcional)
            
        Returns:
            ToolResult con el contenido del archivo
        """
        try:
            # Validar archivo
            validation_result = self.validate_file(file_path)
            if not validation_result.success:
                return validation_result
            
            file_info = validation_result.data
            
            # Verificar que es un archivo de texto
            if not file_info['mime_type'].startswith('text/'):
                return self.create_result(
                    success=False,
                    error=f"Archivo no es de texto: {file_info['mime_type']}"
                )
            
            # Detectar codificación si no se especifica
            if encoding is None and ENCODING_DETECTION_AVAILABLE:
                try:
                    with open(file_path, 'rb') as f:
                        raw_data = f.read(1024)  # Leer primeros 1KB
                        result = chardet.detect(raw_data)
                        encoding = result.get('encoding', 'utf-8')
                except Exception:
                    encoding = 'utf-8'  # Fallback
            elif encoding is None:
                encoding = 'utf-8'  # Fallback por defecto
            
            # Leer archivo con la codificación detectada
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read()
            
            # Estadísticas del archivo
            lines = content.splitlines()
            word_count = len(content.split())
            char_count = len(content)
            
            self.set_completed()
            
            return self.create_result(
                success=True,
                data={
                    'content': content,
                    'encoding': encoding,
                    'lines': len(lines),
                    'words': word_count,
                    'characters': char_count,
                    'file_info': file_info
                },
                file_path=file_path,
                word_count=word_count
            )
            
        except UnicodeDecodeError as e:
            return self.create_result(
                success=False,
                error=f"Error de decodificación: {str(e)}. Intente especificar otra codificación."
            )
        except Exception as e:
            return self.handle_exception(e)
    
    def extract_pdf_text(self, file_path: str) -> ToolResult:
        """
        Extrae texto de un archivo PDF
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            ToolResult con el texto extraído
        """
        if not PDF_AVAILABLE:
            return self.create_result(
                success=False,
                error="PyPDF2 no está instalado. Instale con: pip install PyPDF2"
            )
        
        try:
            # Validar archivo
            validation_result = self.validate_file(file_path)
            if not validation_result.success:
                return validation_result
            
            file_info = validation_result.data
            
            # Verificar que es un PDF
            if file_info['mime_type'] != 'application/pdf':
                return self.create_result(
                    success=False,
                    error="El archivo no es un PDF"
                )
            
            # Extraer texto
            text_content = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- PÁGINA {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        self.logger.warning(f"Error extrayendo página {page_num + 1}: {e}")
                        text_content.append(f"--- PÁGINA {page_num + 1} ---\n[Error extrayendo texto]")
            
            # Combinar todo el texto
            full_text = "\n\n".join(text_content)
            
            # Estadísticas
            lines = full_text.splitlines()
            word_count = len(full_text.split())
            char_count = len(full_text)
            
            self.set_completed()
            
            return self.create_result(
                success=True,
                data={
                    'text': full_text,
                    'pages': page_count,
                    'lines': len(lines),
                    'words': word_count,
                    'characters': char_count,
                    'file_info': file_info
                },
                file_path=file_path,
                page_count=page_count,
                word_count=word_count
            )
            
        except Exception as e:
            return self.handle_exception(e)
    
    def extract_docx_text(self, file_path: str) -> ToolResult:
        """
        Extrae texto de un archivo DOCX
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            ToolResult con el texto extraído
        """
        if not DOCX_AVAILABLE:
            return self.create_result(
                success=False,
                error="python-docx no está instalado. Instale con: pip install python-docx"
            )
        
        try:
            # Validar archivo
            validation_result = self.validate_file(file_path)
            if not validation_result.success:
                return validation_result
            
            file_info = validation_result.data
            
            # Verificar que es un DOCX
            if file_info['mime_type'] != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self.create_result(
                    success=False,
                    error="El archivo no es un DOCX"
                )
            
            # Cargar documento
            doc = Document(file_path)
            
            # Extraer texto de párrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extraer texto de tablas
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combinar todo el texto
            text_parts = []
            if paragraphs:
                text_parts.extend(paragraphs)
            if table_texts:
                text_parts.append("\n--- TABLAS ---")
                text_parts.extend(table_texts)
            
            full_text = "\n".join(text_parts)
            
            # Estadísticas
            lines = full_text.splitlines()
            word_count = len(full_text.split())
            char_count = len(full_text)
            paragraph_count = len(paragraphs)
            table_count = len(doc.tables)
            
            self.set_completed()
            
            return self.create_result(
                success=True,
                data={
                    'text': full_text,
                    'paragraphs': paragraph_count,
                    'tables': table_count,
                    'lines': len(lines),
                    'words': word_count,
                    'characters': char_count,
                    'file_info': file_info
                },
                file_path=file_path,
                word_count=word_count,
                paragraph_count=paragraph_count
            )
            
        except Exception as e:
            return self.handle_exception(e)
    
    def extract_structured_data(self, file_path: str) -> ToolResult:
        """
        Extrae datos estructurados de un archivo (JSON, CSV, etc.)
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            ToolResult con datos estructurados
        """
        try:
            # Validar archivo
            validation_result = self.validate_file(file_path)
            if not validation_result.success:
                return validation_result
            
            file_info = validation_result.data
            file_ext = file_info['extension']
            
            # Procesar según el tipo de archivo
            if file_ext == '.json':
                return self._process_json_file(file_path)
            elif file_ext == '.csv':
                return self._process_csv_file(file_path)
            elif file_ext == '.html':
                return self._process_html_file(file_path)
            else:
                return self.create_result(
                    success=False,
                    error=f"Formato de archivo no soportado para datos estructurados: {file_ext}"
                )
                
        except Exception as e:
            return self.handle_exception(e)
    
    def _process_json_file(self, file_path: str) -> ToolResult:
        """Procesa un archivo JSON"""
        with open(file_path, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                
                # Información adicional sobre el JSON
                info = {
                    'type': type(data).__name__,
                    'keys_count': len(data) if isinstance(data, dict) else 0,
                    'items_count': len(data) if isinstance(data, (list, tuple)) else 1
                }
                
                self.set_completed()
                return self.create_result(
                    success=True,
                    data={
                        'structured_data': data,
                        'info': info
                    },
                    file_path=file_path,
                    json_parsed=True
                )
                
            except json.JSONDecodeError as e:
                return self.create_result(
                    success=False,
                    error=f"Error decodificando JSON: {str(e)}"
                )
    
    def _process_csv_file(self, file_path: str) -> ToolResult:
        """Procesa un archivo CSV"""
        import csv
        
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            data = list(reader)
            
            # Información adicional
            fieldnames = reader.fieldnames if reader.fieldnames else []
            info = {
                'rows': len(data),
                'columns': len(fieldnames),
                'column_names': fieldnames
            }
            
            self.set_completed()
            return self.create_result(
                success=True,
                data={
                    'structured_data': data,
                    'info': info
                },
                file_path=file_path,
                csv_rows=len(data)
            )
    
    def _process_html_file(self, file_path: str) -> ToolResult:
        """Procesa un archivo HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Análisis básico del HTML
            import re
            
            # Contar elementos comunes
            tag_counts = {
                'paragraphs': len(re.findall(r'<p[^>]*>', content, re.IGNORECASE)),
                'headers': len(re.findall(r'<h[1-6][^>]*>', content, re.IGNORECASE)),
                'links': len(re.findall(r'<a[^>]*href=', content, re.IGNORECASE)),
                'images': len(re.findall(r'<img[^>]*src=', content, re.IGNORECASE))
            }
            
            info = {
                'tag_counts': tag_counts,
                'has_title': bool(re.search(r'<title[^>]*>', content, re.IGNORECASE)),
                'has_meta': bool(re.search(r'<meta[^>]*>', content, re.IGNORECASE))
            }
            
            self.set_completed()
            return self.create_result(
                success=True,
                data={
                    'structured_data': {'content': content},
                    'info': info
                },
                file_path=file_path,
                html_analyzed=True
            )
            
        except Exception as e:
            return self.create_result(
                success=False,
                error=f"Error procesando HTML: {str(e)}"
            )
    
    def process_file(self, file_path: str, operation: str = "auto") -> ToolResult:
        """
        Procesa un archivo automáticamente según su tipo
        
        Args:
            file_path: Ruta al archivo
            operation: Tipo de operación ('auto', 'text', 'pdf', 'docx', 'structured')
            
        Returns:
            ToolResult con el resultado del procesamiento
        """
        try:
            # Validar archivo primero
            validation_result = self.validate_file(file_path)
            if not validation_result.success:
                return validation_result
            
            file_info = validation_result.data
            file_ext = file_info['extension']
            mime_type = file_info['mime_type']
            
            # Determinar tipo de operación
            if operation == "auto":
                if mime_type == 'application/pdf':
                    operation = "pdf"
                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    operation = "docx"
                elif mime_type.startswith('text/'):
                    operation = "text"
                elif mime_type in ['application/json', 'text/csv', 'text/html']:
                    operation = "structured"
                else:
                    return self.create_result(
                        success=False,
                        error=f"Tipo de archivo no soportado: {mime_type}"
                    )
            
            # Ejecutar operación específica
            if operation == "text":
                return self.read_text_file(file_path)
            elif operation == "pdf":
                return self.extract_pdf_text(file_path)
            elif operation == "docx":
                return self.extract_docx_text(file_path)
            elif operation == "structured":
                return self.extract_structured_data(file_path)
            else:
                return self.create_result(
                    success=False,
                    error=f"Operación no reconocida: {operation}"
                )
                
        except Exception as e:
            return self.handle_exception(e)
    
    def execute(self, **kwargs) -> ToolResult:
        """
        Ejecuta la funcionalidad principal del procesador de archivos
        
        Args:
            **kwargs: Argumentos que pueden incluir:
                - file_path: Ruta al archivo
                - operation: Tipo de operación
                - encoding: Codificación (para archivos de texto)
                
        Returns:
            ToolResult con el resultado de la operación
        """
        file_path = kwargs.get('file_path')
        if not file_path:
            return self.create_result(
                success=False,
                error="Ruta de archivo requerida"
            )
        
        operation = kwargs.get('operation', 'auto')
        encoding = kwargs.get('encoding')
        
        if operation in ['text', 'read']:
            return self.read_text_file(file_path, encoding)
        elif operation == 'pdf':
            return self.extract_pdf_text(file_path)
        elif operation == 'docx':
            return self.extract_docx_text(file_path)
        elif operation == 'structured':
            return self.extract_structured_data(file_path)
        elif operation == 'validate':
            return self.validate_file(file_path)
        else:
            return self.process_file(file_path, operation)